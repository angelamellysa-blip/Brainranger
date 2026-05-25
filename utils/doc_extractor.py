"""
doc_extractor.py — Extract teks dari PDF dan DOCX untuk BrainRanger.

Digunakan sebagai alternatif input selain foto buku.
PDF scanned (tidak ada teks) akan dideteksi dan dikembalikan sebagai string kosong.
"""

MAX_PAGES  = 10      # Maksimal halaman PDF yang diproses
MAX_WORDS  = 6000    # Maksimal kata (potong kalau lebih)

# ── PDF ──────────────────────────────────────────────────
def extract_pdf(file_bytes: bytes) -> str:
    """
    Ekstrak teks dari PDF.
    Return string teks, atau "" jika PDF scanned/tidak ada teks.
    """
    try:
        import fitz  # pymupdf
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        total_pages = len(doc)
        pages_to_read = min(total_pages, MAX_PAGES)

        texts = []
        for i in range(pages_to_read):
            page = doc[i]
            text = page.get_text("text").strip()
            if text:
                texts.append(f"[Halaman {i + 1}]\n{text}")

        doc.close()

        combined = "\n\n".join(texts)
        return combined
    except Exception as e:
        print(f"extract_pdf error: {e}")
        return ""


def is_scanned_pdf(text: str) -> bool:
    """
    Return True jika PDF tidak punya teks yang cukup (kemungkinan scanned).
    Threshold: < 100 karakter non-whitespace.
    """
    return len(text.replace(" ", "").replace("\n", "")) < 100


# ── DOCX ─────────────────────────────────────────────────
def extract_docx(file_bytes: bytes) -> str:
    """
    Ekstrak teks dari DOCX (paragraf + tabel).
    Return string teks, atau "" jika gagal.
    """
    try:
        import io
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(io.BytesIO(file_bytes))
        parts = []

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # Paragraf biasa
                text = "".join(run.text for run in element.iterchildren(qn("w:r"))
                               if hasattr(run, "text") and run.text)
                # fallback: ambil semua text node
                if not text:
                    from docx.oxml.ns import nsmap
                    text = element.text_content() if hasattr(element, "text_content") else ""
                if text.strip():
                    parts.append(text.strip())

            elif tag == "tbl":
                # Tabel: ambil semua sel
                rows = []
                for row in element.iterchildren(qn("w:tr")):
                    cells = []
                    for cell in row.iterchildren(qn("w:tc")):
                        cell_text = "".join(
                            p.text for p in cell.iter()
                            if hasattr(p, "text") and p.text
                        )
                        cells.append(cell_text.strip())
                    if any(cells):
                        rows.append(" | ".join(cells))
                if rows:
                    parts.append("\n".join(rows))

        return "\n".join(parts)
    except Exception as e:
        print(f"extract_docx error: {e}")
        return ""


# ── Fallback DOCX sederhana ───────────────────────────────
def extract_docx_simple(file_bytes: bytes) -> str:
    """
    Fallback extractor DOCX yang lebih simpel — hanya paragraf, tanpa tabel.
    Dipakai jika extract_docx gagal.
    """
    try:
        import io
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception as e:
        print(f"extract_docx_simple error: {e}")
        return ""


# ── Truncate ─────────────────────────────────────────────
def truncate_content(text: str, max_words: int = MAX_WORDS) -> tuple[str, bool]:
    """
    Potong teks jika melebihi max_words.
    Return: (teks_hasil, was_truncated)
    """
    words = text.split()
    if len(words) <= max_words:
        return text, False
    truncated = " ".join(words[:max_words])
    return truncated, True


# ── Main entry point ──────────────────────────────────────
def extract_document(file_bytes: bytes, mime_type: str) -> dict:
    """
    Entry point utama. Deteksi tipe file, ekstrak teks, truncate jika perlu.

    Return dict:
    {
        "text": str,          # teks hasil ekstrak (sudah di-truncate jika perlu)
        "success": bool,      # True jika ada teks
        "scanned": bool,      # True jika PDF scanned (tidak ada teks)
        "truncated": bool,    # True jika dipotong karena terlalu panjang
        "error": str | None,  # pesan error jika ada
    }
    """
    text = ""
    scanned = False
    error = None

    try:
        if mime_type == "application/pdf":
            text = extract_pdf(file_bytes)
            if is_scanned_pdf(text):
                scanned = True
                text = ""
        elif mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ):
            text = extract_docx(file_bytes)
            if not text.strip():
                # Coba fallback
                text = extract_docx_simple(file_bytes)
        else:
            error = f"Tipe file tidak didukung: {mime_type}"
    except Exception as e:
        error = str(e)
        print(f"extract_document error: {e}")

    text, truncated = truncate_content(text.strip()) if text.strip() else ("", False)

    return {
        "text":      text,
        "success":   bool(text.strip()),
        "scanned":   scanned,
        "truncated": truncated,
        "error":     error,
    }
